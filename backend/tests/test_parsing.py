from unittest.mock import MagicMock, patch

import openpyxl
from docx import Document

from app.core.parsing import parse_document


def test_parse_docx_joins_paragraphs(tmp_path) -> None:
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")
    doc.add_paragraph("Second paragraph.")
    file_path = tmp_path / "sample.docx"
    doc.save(file_path)

    pages = parse_document(file_path, "sample.docx")

    assert len(pages) == 1
    assert pages[0].page is None
    assert "First paragraph." in pages[0].text
    assert "Second paragraph." in pages[0].text


def test_parse_xlsx_one_page_per_sheet(tmp_path) -> None:
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1.append(["name", "value"])
    ws1.append(["alpha", 1])
    ws2 = wb.create_sheet("Sheet2")
    ws2.append(["x", "y"])
    ws2.append(["beta", 2])
    file_path = tmp_path / "sample.xlsx"
    wb.save(file_path)

    pages = parse_document(file_path, "sample.xlsx")

    assert {p.page for p in pages} == {"Sheet1", "Sheet2"}
    sheet1 = next(p for p in pages if p.page == "Sheet1")
    assert "alpha" in sheet1.text


def test_parse_pdf_returns_one_entry_per_nonempty_page() -> None:
    page1 = MagicMock()
    page1.extract_text.return_value = "Page one text"
    page2 = MagicMock()
    page2.extract_text.return_value = "   "  # blank page, should be skipped
    page3 = MagicMock()
    page3.extract_text.return_value = "Page three text"

    mock_pdf = MagicMock()
    mock_pdf.pages = [page1, page2, page3]
    mock_pdf.__enter__.return_value = mock_pdf

    with patch("app.core.parsing.pdfplumber.open", return_value=mock_pdf):
        pages = parse_document("fake.pdf", "fake.pdf")

    assert [(p.page, p.text) for p in pages] == [(1, "Page one text"), (3, "Page three text")]


def test_unsupported_extension_raises() -> None:
    import pytest

    with pytest.raises(ValueError):
        parse_document("file.txt", "file.txt")
