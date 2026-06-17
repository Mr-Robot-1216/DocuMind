"""Generates a sample DOCX used for README screenshots."""
from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "docs" / "sample" / "Acme_Handbook_2025.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
doc.add_heading("Acme Robotics — Employee Handbook 2025", level=0)

doc.add_heading("1. Paid Time Off", level=1)
doc.add_paragraph(
    "New full-time employees accrue 18 days of paid vacation per year during "
    "their first two years of employment. After completing two years of "
    "service, the annual allowance increases to 24 days. Unused vacation of up "
    "to 5 days may be carried over into the following calendar year."
)

doc.add_heading("2. Remote Work", level=1)
doc.add_paragraph(
    "Acme operates on a hybrid model. Employees are expected on-site Tuesday "
    "through Thursday, and may work remotely on Mondays and Fridays. Fully "
    "remote arrangements require director-level approval and are reviewed "
    "every six months."
)

doc.add_heading("3. Equipment and Reimbursement", level=1)
doc.add_paragraph(
    "Each employee receives a laptop and a one-time home-office stipend of "
    "$750. Internet costs are reimbursed up to $40 per month for remote work "
    "days. Reimbursement claims must be submitted within 30 days through the "
    "finance portal."
)

doc.add_heading("4. Professional Development", level=1)
doc.add_paragraph(
    "Acme funds up to $2,000 per employee per year for conferences, courses, "
    "and certifications relevant to the employee's role. Requests are approved "
    "by the employee's manager and must be booked at least four weeks in "
    "advance."
)

doc.save(OUT)
print(f"Wrote {OUT}")
